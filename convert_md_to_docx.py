#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_markdown_file(filepath):
    """마크다운 파일을 읽습니다."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

def convert_markdown_to_docx(md_content, output_path):
    """마크다운을 DOCX로 변환합니다."""
    doc = Document()
    
    # 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # 줄 단위로 처리
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 빈 줄 처리
        if not line.strip():
            i += 1
            continue
        
        # 제목 처리 (# ## ### 등)
        if line.startswith('# '):
            heading_text = line.lstrip('# ').strip()
            heading = doc.add_heading(heading_text, level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        elif line.startswith('## '):
            heading_text = line.lstrip('## ').strip()
            heading = doc.add_heading(heading_text, level=2)
            heading.paragraph_format.space_before = Pt(10)
            heading.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        elif line.startswith('### '):
            heading_text = line.lstrip('### ').strip()
            heading = doc.add_heading(heading_text, level=3)
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after = Pt(4)
            i += 1
            continue
        
        # 표 처리 (마크다운 테이블)
        elif line.strip().startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            # 테이블 헤더
            header_cells = [cell.strip() for cell in line.split('|')[1:-1]]
            
            # 구분선 스킵
            i += 2
            
            # 테이블 바디 수집
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_cells = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                table_rows.append(row_cells)
                i += 1
            
            # 테이블 생성
            if header_cells and table_rows:
                table = doc.add_table(rows=len(table_rows) + 1, cols=len(header_cells))
                table.style = 'Light Grid Accent 1'
                
                # 헤더 행
                hdr_cells = table.rows[0].cells
                for idx, cell_text in enumerate(header_cells):
                    if idx < len(hdr_cells):
                        hdr_cells[idx].text = cell_text
                        # 헤더 포맷
                        for paragraph in hdr_cells[idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                
                # 바디 행
                for row_idx, row in enumerate(table_rows):
                    cells = table.rows[row_idx + 1].cells
                    for col_idx, cell_text in enumerate(row):
                        if col_idx < len(cells):
                            cells[col_idx].text = cell_text
            
            continue
        
        # 이미지 처리
        elif line.strip().startswith('!['):
            match = re.search(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
            if match:
                img_description = match.group(1)
                img_path = match.group(2)
                
                # 이미지 파일이 존재하는지 확인
                img_full_path = Path('/workspaces/icb_slide3') / img_path
                if img_full_path.exists():
                    try:
                        doc.add_picture(str(img_full_path), width=Inches(5.5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        # 이미지 로드 실패 시 텍스트로 표시
                        p = doc.add_paragraph(f"[이미지: {img_description}]")
                        p.paragraph_format.left_indent = Inches(0.25)
                else:
                    p = doc.add_paragraph(f"[이미지: {img_description}]")
                    p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue
        
        # 불릿 리스트 처리 (- 또는 *)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            bullet_text = line.lstrip('- *').strip()
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            i += 1
            
            # 하위 불릿 처리
            while i < len(lines) and (lines[i].startswith('  - ') or lines[i].startswith('  * ')):
                sub_bullet = lines[i].lstrip().lstrip('- *').strip()
                p = doc.add_paragraph(sub_bullet, style='List Bullet 2')
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                i += 1
            continue
        
        # 구분선 처리
        elif line.strip().startswith('---'):
            doc.add_paragraph('_' * 40)
            i += 1
            continue
        
        # 일반 텍스트
        else:
            # 굵은 텍스트와 기울임 처리
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            
            # 간단한 마크다운 포맷 처리
            text = line
            
            # **bold** 처리
            bold_pattern = r'\*\*([^*]+)\*\*'
            parts = re.split(bold_pattern, text)
            
            for idx, part in enumerate(parts):
                if idx % 2 == 0:  # 일반 텍스트
                    # *italic* 처리
                    italic_pattern = r'\*([^*]+)\*'
                    italic_parts = re.split(italic_pattern, part)
                    
                    for jdx, ipart in enumerate(italic_parts):
                        if jdx % 2 == 0:  # 일반 텍스트
                            if ipart:
                                p.add_run(ipart)
                        else:  # 기울임
                            if ipart:
                                run = p.add_run(ipart)
                                run.italic = True
                else:  # 굵은 텍스트
                    run = p.add_run(part)
                    run.bold = True
            
            i += 1
    
    # 문서 저장
    doc.save(output_path)
    print(f"✅ DOCX 파일이 생성되었습니다: {output_path}")

if __name__ == '__main__':
    md_file = '/workspaces/icb_slide3/penguin_analysis_report.md'
    docx_file = '/workspaces/icb_slide3/penguin_analysis_report.docx'
    
    print(f"📄 마크다운 파일 변환 시작...")
    print(f"입력: {md_file}")
    print(f"출력: {docx_file}")
    
    md_content = read_markdown_file(md_file)
    convert_markdown_to_docx(md_content, docx_file)
    
    print(f"✅ 변환 완료!")
